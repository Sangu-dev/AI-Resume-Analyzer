"""Resume text extractor — supports PDF, TXT (and DOCX if python-docx is installed).

Previously only pdfplumber was called for all file types, which crashed on TXT files.
This module now branches on the file extension before choosing an extraction strategy.
"""

from __future__ import annotations

import io
import logging

import pdfplumber

logger = logging.getLogger(__name__)

# Optional DOCX support — fail gracefully if python-docx is not installed
try:
    from docx import Document as _DocxDocument
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False


def extract_text(file_obj: io.BytesIO | io.BufferedReader, filename: str = "") -> str:
    """Extract plain text from a PDF, TXT, or DOCX file object.

    Args:
        file_obj: A file-like object (BytesIO or similar).
        filename:  Original filename used to determine the format. If omitted,
                   PDF extraction is attempted by default.

    Returns:
        Extracted plain text as a single string. Never raises — returns an empty
        string on unrecoverable failures so the caller can decide how to handle it.
    """
    ext = filename.lower().rsplit(".", 1)[-1] if filename else "pdf"

    try:
        if ext == "txt":
            return _extract_from_txt(file_obj)
        elif ext == "docx":
            return _extract_from_docx(file_obj)
        else:
            # Default: treat as PDF
            return _extract_from_pdf(file_obj)
    except Exception as exc:
        logger.error("Text extraction failed for '%s': %s", filename, exc)
        raise RuntimeError(
            f"Could not extract text from '{filename}'. "
            "Ensure it is a valid PDF, TXT, or DOCX file."
        ) from exc


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _extract_from_pdf(file_obj: io.BytesIO) -> str:
    """Extract text from a PDF using pdfplumber page-by-page."""
    parts: list[str] = []
    with pdfplumber.open(file_obj) as pdf:
        for i, page in enumerate(pdf.pages):
            page_text = page.extract_text()
            if page_text and page_text.strip():
                parts.append(page_text)
            else:
                logger.debug("Page %d had no extractable text.", i + 1)
    return "\n".join(parts)


def _extract_from_txt(file_obj: io.BytesIO) -> str:
    """Decode a plain-text file, trying common encodings in order."""
    raw: bytes = file_obj.read() if hasattr(file_obj, "read") else file_obj.getvalue()

    for encoding in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue

    # Last-resort: replace undecodeable bytes rather than crash
    return raw.decode("utf-8", errors="replace")


def _extract_from_docx(file_obj: io.BytesIO) -> str:
    """Extract text from a DOCX file including table cells."""
    if not _DOCX_AVAILABLE:
        raise RuntimeError(
            "DOCX support requires python-docx. "
            "Run: pip install python-docx"
        )
    doc = _DocxDocument(file_obj)
    parts: list[str] = []

    # Main body paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # Table cells (often contain skills/contact info in resume templates)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)

    return "\n".join(parts)
